import os
import re
import asyncio
import warnings
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.utils.keyboard import InlineKeyboardBuilder
from g4f.client import Client
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import uuid

warnings.filterwarnings("ignore", category=RuntimeWarning)

# Конфигурация
ADMIN_ID = 1404494933  # ID админа
ONE_TIME_KEYS = set()  # Множество для хранения одноразовых ключей

class Form(StatesGroup):
    theme = State()
    working = State()
    access_code = State()

client = Client()
bot = Bot(token="7722309470:AAF_WkaBZvK5iyS3kw0O2dpfLYrGLmG5o5k")
dp = Dispatcher()

def sanitize_filename(name):
    clean_name = re.sub(r'[\\/*?:"<>|]', "", name)
    clean_name = clean_name.replace(" ", "_")[:50]
    return clean_name.strip()

def get_filename(user_id, theme):
    safe_theme = sanitize_filename(theme)
    return f"{safe_theme}_{user_id}.md"

def format_response(text):
    text = re.sub(r'```(\w+)?\n(.*?)```', lambda m: f"\n{m.group(2)}\n", text, flags=re.DOTALL)
    text = re.sub(r'`(.*?)`', lambda m: f"{m.group(1)}", text)
    return text.strip()

async def generate_structure(theme):
    try:
        prompt = (
            f"Составь подробную структуру диплома на тему '{theme}' с нумерацией разделов по шаблону X.X. "
            "Включи 4 главы с 3 подразделами каждая. Пример формата: "
            "4. Глава 4\n4.1. Подраздел 1\n4.2. Подраздел 2"
        )
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            web_search=False
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Ошибка генерации структуры: {str(e)}"

async def generate_subsection_content(theme, section):
    try:
        prompt = (
            f"Напиши академический текст для подраздела '{section}' диплома на тему '{theme}'. "
            "Объем: 3000-3500 слов обязательно. Используй подзаголовки, списки и примеры. "
            "Оформи как профессиональный научный текст без markdown."
        )
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            web_search=False
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Ошибка генерации раздела: {str(e)}"

def update_md_file(user_id, theme, content, section):
    filename = get_filename(user_id, theme)
    mode = "a" if os.path.exists(filename) else "w"
    
    level = section.count('.') + 1
    hashes = '#' * min(level + 1, 4)
    
    with open(filename, mode, encoding="utf-8") as f:
        f.write(f"\n{hashes} {section}\n\n")
        f.write(content + "\n")
    return filename

def add_heading_with_bookmark(doc, text, level, bookmark_name):
    heading = doc.add_heading(text, level)
    paragraph = heading._element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)
    paragraph.append(bookmark_start)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    bookmark_end.set(qn('w:name'), bookmark_name)
    paragraph.append(bookmark_end)
    return heading

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "TOC \\o '1-3' \\h \\z \\u"
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def md_to_docx(md_filename, docx_filename):
    doc = Document()
    doc.add_heading("Оглавление", level=1)
    add_table_of_contents(doc)

    with open(md_filename, "r", encoding="utf-8") as md_file:
        lines = md_file.readlines()

    for line in lines:
        line = line.strip()

        if line.startswith("#"):
            header_level = line.count("#")
            header_text = line.lstrip("#").strip()
            bookmark_name = f"heading_{header_text.replace(' ', '_')}"
            add_heading_with_bookmark(doc, header_text, min(header_level, 6), bookmark_name)
        elif "**" in line:
            bold_texts = re.findall(r"\*\*(.*?)\*\*", line)
            paragraph = doc.add_paragraph()
            parts = re.split(r"\*\*.*?\*\*", line)
            for i, part in enumerate(parts):
                paragraph.add_run(part)
                if i < len(bold_texts):
                    run = paragraph.add_run(bold_texts[i])
                    run.bold = True
        elif "*" in line:
            italic_texts = re.findall(r"\*(.*?)\*", line)
            paragraph = doc.add_paragraph()
            parts = re.split(r"\*.*?\*", line)
            for i, part in enumerate(parts):
                paragraph.add_run(part)
                if i < len(italic_texts):
                    run = paragraph.add_run(italic_texts[i])
                    run.italic = True
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(line.lstrip("- ").strip())
        elif re.match(r"^\d+\. ", line):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(re.sub(r"^\d+\. ", "", line).strip())
        else:
            doc.add_paragraph(line)

    doc.save(docx_filename)
    print(f"Файл '{docx_filename}' успешно создан с оглавлением и гиперссылками!")

async def delete_files(*filenames):
    for filename in filenames:
        if os.path.exists(filename):
            os.remove(filename)
            print(f"Файл '{filename}' удален.")

@dp.message(Command("start"))
async def start_handler(msg: types.Message, state: FSMContext):
    if msg.from_user.id == ADMIN_ID:
        await msg.answer("👋 Привет, админ! Используй /generate_key для создания ключа доступа.")
    else:
        await msg.answer("🔑 Введите код доступа для начала работы:")
        await state.set_state(Form.access_code)

@dp.message(Form.access_code)
async def process_access_code(msg: types.Message, state: FSMContext):
    access_code = msg.text.strip()
    if access_code in ONE_TIME_KEYS:
        ONE_TIME_KEYS.remove(access_code)
        await msg.answer("✅ Код доступа принят. Введите тему дипломной работы:")
        await state.set_state(Form.theme)
    else:
        await msg.answer("❌ Неверный код доступа. Обратитесь к @TETRIX_UNO.")

@dp.message(Command("generate_key"))
async def generate_key(msg: types.Message):
    if msg.from_user.id == ADMIN_ID:
        key = str(uuid.uuid4())
        ONE_TIME_KEYS.add(key)
        await msg.answer(f"🔑 Одноразовый ключ доступа: {key}")
    else:
        await msg.answer("❌ У вас нет прав для выполнения этой команды.")

@dp.message(Form.theme)
async def process_theme(msg: types.Message, state: FSMContext):
    theme = msg.text.strip()
    await state.update_data(theme=theme, sections=[], current_section=0)
    
    loading_msg = await msg.answer("🔄 Генерируем структуру...")
    structure = await generate_structure(theme)
    
    sections = re.findall(r'\d+(?:\.\d+)*\..+?(?=\n\d|\Z)', structure, flags=re.DOTALL)
    sections = [s.strip() for s in sections if re.match(r'\d+\.', s)]
    
    await bot.delete_message(msg.chat.id, loading_msg.message_id)
    
    if "Ошибка" in structure or not sections:
        await msg.answer("❌ Ошибка при генерации структуры")
        return
    
    builder = InlineKeyboardBuilder()
    builder.row(
        types.InlineKeyboardButton(text="🔄 Перегенерировать", callback_data="regenerate_structure"),
        types.InlineKeyboardButton(text="✅ Начать генерацию", callback_data="start_generation")
    )
    
    await state.update_data(sections=sections)
    update_md_file(msg.from_user.id, theme, structure, "Структура работы")
    
    await msg.answer(
        f"📚 Сгенерированная структура:\n\n{format_response(structure)}",
        reply_markup=builder.as_markup()
    )

@dp.callback_query(F.data == "regenerate_structure")
async def regenerate_structure(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()
    theme = data['theme']
    
    loading_msg = await callback.message.answer("🔄 Создаем новую структуру...")
    structure = await generate_structure(theme)
    sections = re.findall(r'\d+(?:\.\d+)*\..+?(?=\n\d|\Z)', structure, flags=re.DOTALL)
    sections = [s.strip() for s in sections if re.match(r'\d+\.', s)]
    
    await bot.delete_message(callback.message.chat.id, loading_msg.message_id)
    
    if "Ошибка" in structure or not sections:
        await callback.message.answer("❌ Ошибка при перегенерации структуры")
        return
    
    builder = InlineKeyboardBuilder()
    builder.row(
        types.InlineKeyboardButton(text="🔄 Еще раз", callback_data="regenerate_structure"),
        types.InlineKeyboardButton(text="✅ Начать генерацию", callback_data="start_generation")
    )
    
    await state.update_data(sections=sections)
    update_md_file(callback.from_user.id, theme, structure, "Обновленная структура")
    
    await callback.message.edit_text(
        f"📚 Новая структура:\n\n{format_response(structure)}",
        reply_markup=builder.as_markup()
    )

@dp.callback_query(F.data == "start_generation")
async def start_generation(callback: types.CallbackQuery, state: FSMContext):
    await callback.answer()
    data = await state.get_data()
    await state.set_state(Form.working)
    await process_next_section(callback.message, state, data)

async def process_next_section(msg, state, data):
    sections = data['sections']
    current = data['current_section']
    
    if current >= len(sections):
        filename = get_filename(msg.from_user.id, data['theme'])
        docx_filename = filename.replace(".md", ".docx")
        md_to_docx(filename, docx_filename)
        
        await msg.answer_document(
            types.FSInputFile(filename),
            caption="🎓 Дипломная работа готова!",
            reply_markup=types.ReplyKeyboardRemove()
        )
        await msg.answer_document(
            types.FSInputFile(docx_filename),
            caption="📄 Word-версия документа:",
            reply_markup=types.ReplyKeyboardRemove()
        )
        
        # Удаляем файлы после отправки
        await delete_files(filename, docx_filename)
        
        await state.clear()
        return
    
    section = sections[current].strip()
    loading_msg = await msg.answer(f"⏳ Генерируем раздел {current+1}/{len(sections)}:\n{section[:40]}...")
    
    content = await generate_subsection_content(data['theme'], section)
    update_md_file(msg.from_user.id, data['theme'], content, section)
    
    await bot.delete_message(msg.chat.id, loading_msg.message_id)
    await msg.answer(f"✅ Раздел {current+1}/{len(sections)} готов!\n{section}")
    
    await state.update_data(current_section=current+1)
    new_data = await state.get_data()
    await process_next_section(msg, state, new_data)

@dp.message(Command("get_file"))
async def send_file(msg: types.Message, state: FSMContext):
    data = await state.get_data()
    theme = data.get('theme', 'diploma')
    filename = get_filename(msg.from_user.id, theme)
    
    if os.path.exists(filename):
        await msg.answer_document(
            types.FSInputFile(filename),
            caption=f"📁 Текущая версия работы: {theme}"
        )
    else:
        await msg.answer("❌ Файл еще не создан")

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
